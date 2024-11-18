import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE

def get_multiline_input(prompt):
    print(prompt)
    lines = []
    while True:
        line = input()
        if line == '':
            break
        lines.append(line)
    return '\n'.join(lines)

def parse_personal_info(data):
    personal_info = {}
    lines = data.strip().split('\n')
    personal_info['name'] = lines[0]
    for line in lines[1:]:
        if ':' in line:
            key, value = line.split(':', 1)
            personal_info[key.strip().lower()] = value.strip()
    return personal_info

def get_skills():
    skills = {}
    while True:
        category = input("Enter skill category (or press Enter to finish): ").strip()
        if category == '':
            break
        items = input(f"Enter skills for {category} (comma-separated): ").strip()
        skills[category] = [item.strip() for item in items.split(',')]
    return skills

def get_experiences():
    experiences = []
    while True:
        add_experience = input("Do you want to add an experience? (yes/no): ").strip().lower()
        if add_experience != 'yes':
            break
        position = input("Enter Position Title: ").strip()
        company = input("Enter Company Name: ").strip()
        dates = input("Enter Dates (e.g., 01/2020 - 12/2021): ").strip()
        details = []
        print("Enter details or achievements (Press Enter twice to finish):")
        while True:
            detail = input()
            if detail == '':
                break
            details.append(detail)
        experience = {
            'position': position,
            'company': company,
            'dates': dates,
            'details': details
        }
        experiences.append(experience)
    return experiences

def get_education():
    education_list = []
    while True:
        add_education = input("Do you want to add an education entry? (yes/no): ").strip().lower()
        if add_education != 'yes':
            break
        degree = input("Enter Degree Title: ").strip()
        institution = input("Enter Institution Name: ").strip()
        year = input("Enter Graduation Year: ").strip()
        edu_entry = {
            'degree': degree,
            'institution': institution,
            'year': year
        }
        education_list.append(edu_entry)
    return education_list

def get_list_section(section_name):
    items = []
    print(f"\nEnter {section_name} (Press Enter twice to finish):")
    while True:
        item = input()
        if item == '':
            break
        items.append(item)
    return items

def create_resume(personal_info, skills, experiences, education, certifications, hobbies, languages, personal_details, profile_pic_path):
    document = Document()

    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add a style for the name
    styles = document.styles
    if 'Name Style' not in styles:
        name_style = styles.add_style('Name Style', WD_STYLE_TYPE.PARAGRAPH)
        name_font = name_style.font
        name_font.name = 'Calibri'
        name_font.size = Pt(24)
        name_font.bold = True

    # Create a table for the header (picture and contact info)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(4.5)

    # Add Profile Picture
    if profile_pic_path:
        try:
            cell = table.cell(0, 0)
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(profile_pic_path, width=Inches(1.5))
        except Exception as e:
            print(f"Error adding profile picture: {e}")

    # Add Name and Contact Info
    cell = table.cell(0, 1)
    paragraph = cell.paragraphs[0]
    paragraph.style = document.styles['Name Style']
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph.add_run(personal_info.get('name', ''))

    # Contact Information
    contact_lines = []
    for key in ['location', 'phone', 'email', 'linkedin', 'github']:
        if key in personal_info:
            contact_lines.append(f"{key.capitalize()}: {personal_info[key]}")
    contact_info = ' | '.join(contact_lines)
    p = cell.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run(contact_info)

    # Add a horizontal line
    document.add_paragraph().add_run().add_break()
    document.add_paragraph().add_run('')

    # Add Skills Section
    if skills:
        document.add_heading('Skills', level=1)
        for category, items in skills.items():
            p = document.add_paragraph()
            p.add_run(f"{category}: ").bold = True
            p.add_run(', '.join(items))

    # Add Experience Section
    if experiences:
        document.add_heading('Experience', level=1)
        for exp in experiences:
            p = document.add_paragraph()
            p.add_run(f"{exp['position']}, ").bold = True
            p.add_run(f"{exp['company']} ({exp['dates']})")
            for detail in exp['details']:
                document.add_paragraph(detail.strip(), style='List Bullet')

    # Add Education Section
    if education:
        document.add_heading('Education', level=1)
        for edu in education:
            p = document.add_paragraph()
            p.add_run(f"{edu['degree']}, ").bold = True
            p.add_run(f"{edu['institution']} ({edu['year']})")

    # Add Certifications
    if certifications:
        document.add_heading('Certifications', level=1)
        for cert in certifications:
            document.add_paragraph(cert, style='List Bullet')

    # Add Hobbies and Interests
    if hobbies:
        document.add_heading('Hobbies and Interests', level=1)
        for hobby in hobbies:
            document.add_paragraph(hobby, style='List Bullet')

    # Add Languages
    if languages:
        document.add_heading('Languages', level=1)
        languages_paragraph = document.add_paragraph()
        languages_paragraph.add_run(', '.join(languages))

    # Add Personal Information
    if personal_details:
        document.add_heading('Personal Information', level=1)
        for detail in personal_details:
            document.add_paragraph(detail, style='List Bullet')

    # Save the document
    document.save('Your_Resume1.docx')
    print("Resume created successfully!")

def main():
    print("=== Resume Builder ===\n")

    personal_info_input = get_multiline_input("Enter Personal Information (Press Enter twice to finish):")
    personal_info = parse_personal_info(personal_info_input)

    # Ask for Profile Picture Path
    profile_pic_path = input("\nEnter the file path to your profile picture (or press Enter to skip): ").strip()
    if profile_pic_path == '':
        profile_pic_path = None

    # Skills
    print("\n=== Skills Section ===")
    skills = get_skills()

    # Experience
    print("\n=== Experience Section ===")
    experiences = get_experiences()

    # Education
    print("\n=== Education Section ===")
    education = get_education()

    # Certifications
    certifications = get_list_section("Certifications")

    # Hobbies and Interests
    hobbies = get_list_section("Hobbies and Interests")

    # Languages
    languages = get_list_section("Languages")

    # Personal Details
    personal_details = get_list_section("Personal Information (Date of Birth, Gender, etc.)")

    create_resume(personal_info, skills, experiences, education, certifications, hobbies, languages, personal_details, profile_pic_path)

if __name__ == '__main__':
    main()
